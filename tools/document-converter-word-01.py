from collections.abc import Generator
from typing import Any
import io
import os
import sys

from dify_plugin import Tool
from dify_plugin.entities.tool import ToolInvokeMessage


class DocumentConverterWord01Tool(Tool):
    def _invoke(self, tool_parameters: dict[str, Any]) -> Generator[ToolInvokeMessage]:
        markdown_text = (tool_parameters or {}).get("markdown")
        filename = (tool_parameters or {}).get("filename") or "document.docx"
        style_profile = (tool_parameters or {}).get("style_profile") or "学术论文"
        # Optional overrides
        # H1-H5 font & size overrides
        h1_font = (tool_parameters or {}).get("h1_font")
        h1_size_pt_raw = (tool_parameters or {}).get("h1_size_pt")
        h2_font = (tool_parameters or {}).get("h2_font")
        h2_size_pt_raw = (tool_parameters or {}).get("h2_size_pt")
        h3_font = (tool_parameters or {}).get("h3_font")
        h3_size_pt_raw = (tool_parameters or {}).get("h3_size_pt")
        h4_font = (tool_parameters or {}).get("h4_font")
        h4_size_pt_raw = (tool_parameters or {}).get("h4_size_pt")
        h5_font = (tool_parameters or {}).get("h5_font")
        h5_size_pt_raw = (tool_parameters or {}).get("h5_size_pt")
        # Body overrides
        body_font = (tool_parameters or {}).get("body_font")
        body_size_pt_raw = (tool_parameters or {}).get("body_size_pt")

        def _to_float(val: Any) -> float | None:
            try:
                if val is None:
                    return None
                if isinstance(val, (int, float)):
                    return float(val)
                s = str(val).strip()
                if not s:
                    return None
                return float(s)
            except Exception:
                return None

        h1_size_pt = _to_float(h1_size_pt_raw)
        h2_size_pt = _to_float(h2_size_pt_raw)
        h3_size_pt = _to_float(h3_size_pt_raw)
        h4_size_pt = _to_float(h4_size_pt_raw)
        h5_size_pt = _to_float(h5_size_pt_raw)
        body_size_pt = _to_float(body_size_pt_raw)

        if not markdown_text or not isinstance(markdown_text, str):
            yield self.create_text_message("Parameter 'markdown' is required and must be a string.")
            return

        # Ensure filename ends with .docx
        if not isinstance(filename, str) or not filename.strip():
            filename = "document.docx"
        filename = filename.strip()
        if not filename.lower().endswith(".docx"):
            filename += ".docx"

        try:
            # Convert Markdown -> HTML -> DOCX bytes (base document)
            from markdown import markdown
            from html2docx import html2docx

            html = markdown(markdown_text)
            base_buf = html2docx(html, title=filename)

            # Optionally apply Chinese typography styles using python-docx
            try:
                from docx import Document
                from docx.shared import Inches, Pt, RGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
                from docx.enum.section import WD_ORIENT
                from docx.oxml.ns import qn

                # Try import local style manager from workspace root
                root_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
                if os.path.exists(os.path.join(root_dir, "document_styles.py")) and root_dir not in sys.path:
                    sys.path.insert(0, root_dir)
                try:
                    from document_styles import DocumentStyleManager, STYLE_TEMPLATES, FontStyle, ParagraphStyle
                except Exception:
                    # Fallback: minimal defaults if style manager not importable
                    DocumentStyleManager = None
                    STYLE_TEMPLATES = {}
                    FontStyle = None
                    ParagraphStyle = None
                    from types import SimpleNamespace as _NS

                doc = Document(base_buf)

                # Helper functions to apply styles
                def apply_font_style(run, font_style):
                    if font_style is None:
                        return
                    # English font
                    run.font.name = font_style.family
                    # Chinese East Asian font
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_style.family)
                    # Size & weight
                    run.font.size = Pt(font_style.size)
                    run.font.bold = bool(getattr(font_style, 'bold', False))
                    run.font.italic = bool(getattr(font_style, 'italic', False))
                    # Color
                    color = getattr(font_style, 'color', None)
                    if color and isinstance(color, str) and color.startswith('#') and len(color) == 7:
                        r = int(color[1:3], 16)
                        g = int(color[3:5], 16)
                        b = int(color[5:7], 16)
                        run.font.color.rgb = RGBColor(r, g, b)

                def apply_paragraph_style(paragraph, paragraph_style):
                    if paragraph_style is None:
                        return
                    alignment_map = {
                        "left": WD_ALIGN_PARAGRAPH.LEFT,
                        "center": WD_ALIGN_PARAGRAPH.CENTER,
                        "right": WD_ALIGN_PARAGRAPH.RIGHT,
                        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
                    }
                    paragraph.alignment = alignment_map.get(getattr(paragraph_style, 'alignment', 'left'), WD_ALIGN_PARAGRAPH.LEFT)
                    # Line spacing (multiple)
                    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                    ls = getattr(paragraph_style, 'line_spacing', 1.5)
                    paragraph.paragraph_format.line_spacing = ls
                    # Space before/after
                    sb = getattr(paragraph_style, 'space_before', 0)
                    sa = getattr(paragraph_style, 'space_after', 0)
                    paragraph.paragraph_format.space_before = Pt(sb)
                    paragraph.paragraph_format.space_after = Pt(sa)
                    # Indents
                    if getattr(paragraph_style, 'indent_first_line', 0) > 0:
                        paragraph.paragraph_format.first_line_indent = Pt(paragraph_style.indent_first_line)
                    if getattr(paragraph_style, 'indent_left', 0) > 0:
                        paragraph.paragraph_format.left_indent = Pt(paragraph_style.indent_left)
                    if getattr(paragraph_style, 'indent_right', 0) > 0:
                        paragraph.paragraph_format.right_indent = Pt(paragraph_style.indent_right)

                def apply_page_style(document, page_style):
                    if page_style is None:
                        return
                    section = document.sections[0]
                    # mm to inches
                    def mm_to_inches(mm):
                        return mm / 25.4
                    section.page_width = Inches(mm_to_inches(getattr(page_style, 'width', 210)))
                    section.page_height = Inches(mm_to_inches(getattr(page_style, 'height', 297)))
                    section.top_margin = Inches(mm_to_inches(getattr(page_style, 'margin_top', 25)))
                    section.bottom_margin = Inches(mm_to_inches(getattr(page_style, 'margin_bottom', 25)))
                    section.left_margin = Inches(mm_to_inches(getattr(page_style, 'margin_left', 25)))
                    section.right_margin = Inches(mm_to_inches(getattr(page_style, 'margin_right', 25)))
                    if getattr(page_style, 'orientation', 'portrait') == 'landscape':
                        section.orientation = WD_ORIENT.LANDSCAPE

                # Build style manager and apply template
                sm = DocumentStyleManager() if DocumentStyleManager else None
                if sm and STYLE_TEMPLATES:
                    tpl = STYLE_TEMPLATES.get(style_profile)
                    if tpl:
                        # update known styles
                        if 'normal' in tpl:
                            sm.update_style('normal', tpl['normal'].get('font'), tpl['normal'].get('paragraph'))
                        if 'heading1' in tpl:
                            sm.update_style('heading1', tpl['heading1'].get('font'), tpl['heading1'].get('paragraph'))
                        if 'heading2' in tpl:
                            sm.update_style('heading2', tpl['heading2'].get('font'), tpl['heading2'].get('paragraph'))
                        if 'heading3' in tpl:
                            sm.update_style('heading3', tpl['heading3'].get('font'), tpl['heading3'].get('paragraph'))
                        if 'heading4' in tpl:
                            sm.update_style('heading4', tpl['heading4'].get('font'), tpl['heading4'].get('paragraph'))
                        if 'heading5' in tpl:
                            sm.update_style('heading5', tpl['heading5'].get('font'), tpl['heading5'].get('paragraph'))
                        if 'title' in tpl:
                            sm.update_style('title', tpl['title'].get('font'), tpl['title'].get('paragraph'))
                        # page style
                        if 'page' in tpl:
                            sm.set_page_style(tpl['page'])

                # Apply overrides on top of template
                if sm:
                    # Headings overrides H1..H5
                    heading_overrides = [
                        ("heading1", h1_font, h1_size_pt),
                        ("heading2", h2_font, h2_size_pt),
                        ("heading3", h3_font, h3_size_pt),
                        ("heading4", h4_font, h4_size_pt),
                        ("heading5", h5_font, h5_size_pt),
                    ]
                    for key, font_name, size_pt in heading_overrides:
                        base = sm.get_style(key)
                        base_font = base.get('font')
                        base_para = base.get('paragraph')
                        new_font = base_font
                        new_para = base_para
                        if FontStyle:
                            if font_name or size_pt is not None:
                                new_font = FontStyle(
                                    family=font_name or getattr(base_font, 'family', '宋体'),
                                    size=int(size_pt) if size_pt is not None else getattr(base_font, 'size', 12),
                                    bold=True,
                                    italic=getattr(base_font, 'italic', False),
                                    color="#000000",
                                )
                        if ParagraphStyle:
                            new_para = ParagraphStyle(
                                alignment=('center' if key == 'heading1' else 'left'),
                                line_spacing=1.0,
                                space_before=getattr(base_para, 'space_before', 0),
                                space_after=getattr(base_para, 'space_after', 0),
                                indent_first_line=getattr(base_para, 'indent_first_line', 0),
                                indent_left=getattr(base_para, 'indent_left', 0),
                                indent_right=getattr(base_para, 'indent_right', 0),
                            )
                        sm.update_style(key, new_font if (font_name or size_pt is not None) else None, new_para)

                    # Body overrides (Normal)
                    if body_font or body_size_pt is not None:
                        base = sm.get_style('normal')
                        base_font = base.get('font')
                        base_para = base.get('paragraph')
                        new_font = base_font
                        new_para = base_para
                        if FontStyle:
                            new_font = FontStyle(
                                family=body_font or getattr(base_font, 'family', '宋体'),
                                size=int(body_size_pt) if body_size_pt is not None else getattr(base_font, 'size', 12),
                                bold=getattr(base_font, 'bold', False),
                                italic=getattr(base_font, 'italic', False),
                                color="#000000",
                            )
                        if ParagraphStyle:
                            new_para = ParagraphStyle(
                                alignment=getattr(base_para, 'alignment', 'left'),
                                line_spacing=1.0,
                                space_before=getattr(base_para, 'space_before', 0),
                                space_after=getattr(base_para, 'space_after', 0),
                                indent_first_line=getattr(base_para, 'indent_first_line', 24),
                                indent_left=getattr(base_para, 'indent_left', 0),
                                indent_right=getattr(base_para, 'indent_right', 0),
                            )
                        sm.update_style('normal', new_font, new_para)

                # Apply page style first
                if sm:
                    apply_page_style(doc, getattr(sm, 'page_style', None))

                # Map docx built-in styles to our manager keys
                style_map = {
                    'Normal': 'normal',
                    'Heading 1': 'heading1',
                    'Heading 2': 'heading2',
                    'Heading 3': 'heading3',
                    'Heading 4': 'heading4',
                    'Heading 5': 'heading5',
                    'Heading 6': 'heading5',
                    'List Paragraph': 'normal',
                    'Quote': 'quote',
                    'Intense Quote': 'quote',
                    'Preformatted': 'code',
                    'Code': 'code',
                }

                # Apply styles to paragraphs and runs
                if sm:
                    for p in doc.paragraphs:
                        key = style_map.get(getattr(p.style, 'name', 'Normal'), 'normal')
                        style_def = sm.get_style(key)
                        # Runs font
                        for r in p.runs:
                            apply_font_style(r, style_def.get('font'))
                        # Paragraph formatting
                        apply_paragraph_style(p, style_def.get('paragraph'))
                else:
                    # No style manager: apply simple overrides per heading level
                    from types import SimpleNamespace as _NS
                    # Build simple font objects
                    h_fonts = {
                        'Heading 1': _NS(family=h1_font or '宋体', size=int(h1_size_pt) if h1_size_pt is not None else 16, bold=True, italic=False, color="#000000"),
                        'Heading 2': _NS(family=h2_font or '宋体', size=int(h2_size_pt) if h2_size_pt is not None else 14, bold=True, italic=False, color="#000000"),
                        'Heading 3': _NS(family=h3_font or '宋体', size=int(h3_size_pt) if h3_size_pt is not None else 12, bold=True, italic=False, color="#000000"),
                        'Heading 4': _NS(family=h4_font or '宋体', size=int(h4_size_pt) if h4_size_pt is not None else 11, bold=True, italic=False, color="#000000"),
                        'Heading 5': _NS(family=h5_font or '宋体', size=int(h5_size_pt) if h5_size_pt is not None else 10, bold=True, italic=False, color="#000000"),
                    }
                    body_font_obj = _NS(family=body_font or '宋体', size=int(body_size_pt) if body_size_pt is not None else 12, bold=False, italic=False, color="#000000")
                    heading_para_obj_h1 = _NS(alignment='center', line_spacing=1.0, space_before=0, space_after=0, indent_first_line=0, indent_left=0, indent_right=0)
                    heading_para_obj_other = _NS(alignment='left', line_spacing=1.0, space_before=0, space_after=0, indent_first_line=0, indent_left=0, indent_right=0)
                    body_para_obj = _NS(alignment='left', line_spacing=1.0, space_before=0, space_after=0, indent_first_line=24, indent_left=0, indent_right=0)
                    for p in doc.paragraphs:
                        name = getattr(p.style, 'name', 'Normal')
                        if name.startswith('Heading'):
                            font_obj = h_fonts.get(name, h_fonts['Heading 3'])
                            for r in p.runs:
                                apply_font_style(r, font_obj)
                            if name == 'Heading 1':
                                apply_paragraph_style(p, heading_para_obj_h1)
                            else:
                                apply_paragraph_style(p, heading_para_obj_other)
                        else:
                            for r in p.runs:
                                apply_font_style(r, body_font_obj)
                            apply_paragraph_style(p, body_para_obj)

                # Save to bytes
                out_buf = io.BytesIO()
                doc.save(out_buf)
                docx_bytes = out_buf.getvalue()
            except Exception:
                # If styling fails, fall back to base output
                docx_bytes = base_buf.getvalue()

            # Return as file blob with proper mime type
            yield self.create_blob_message(
                blob=docx_bytes,
                meta={
                    "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "filename": filename,
                },
            )

            # Also return a small JSON summary for workflow consumption
            yield self.create_json_message(
                {
                    "filename": filename,
                    "size_bytes": len(docx_bytes),
                    "style_profile": style_profile,
                    "overrides": {
                        "h1_font": h1_font,
                        "h1_size_pt": h1_size_pt,
                        "h2_font": h2_font,
                        "h2_size_pt": h2_size_pt,
                        "h3_font": h3_font,
                        "h3_size_pt": h3_size_pt,
                        "h4_font": h4_font,
                        "h4_size_pt": h4_size_pt,
                        "h5_font": h5_font,
                        "h5_size_pt": h5_size_pt,
                        "body_font": body_font,
                        "body_size_pt": body_size_pt,
                    },
                    "message": "Markdown converted to Word with Chinese typography",
                }
            )
        except Exception as e:
            # Surface the error to users
            yield self.create_text_message(f"Conversion failed: {e}")
